# Copyright (c) Opendatalab. All rights reserved.
"""
Regression test for docx page-break detection.

DocxConverter's page splitting previously only fired on Word *section*
breaks (<w:sectPr>) — manual page breaks (<w:br w:type="page"/>, inserted via
Word's "Insert > Page Break" / Ctrl+Enter) and the pageBreakBefore paragraph
property were never inspected, so any docx built with manual breaks instead
of section breaks collapsed to a single page regardless of its real length.
"""
from io import BytesIO

from docx import Document

from mineru.model.docx.main import convert_binary


def _build_repro_docx() -> bytes:
    doc = Document()
    doc.add_paragraph("Page 0 content")

    for i in range(1, 4):
        doc.add_page_break()  # <w:br w:type="page"/>
        doc.add_paragraph(f"Page {i} content (manual page break)")

    p = doc.add_paragraph("Page 4 content (pageBreakBefore)")
    p.paragraph_format.page_break_before = True

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()


def test_manual_page_breaks_and_page_break_before_split_pages():
    docx_bytes = _build_repro_docx()
    pages = convert_binary(BytesIO(docx_bytes))

    # 1 initial page + 3 manual <w:br w:type="page"/> breaks + 1
    # pageBreakBefore paragraph = 5 pages, even though this document has
    # zero <w:sectPr> section breaks.
    assert len(pages) == 5


def test_disabled_page_break_before_does_not_split():
    doc = Document()
    doc.add_paragraph("Page 0 content")
    p = doc.add_paragraph("Still page 0 — property explicitly disabled")
    p.paragraph_format.page_break_before = False

    buf = BytesIO()
    doc.save(buf)
    pages = convert_binary(BytesIO(buf.getvalue()))

    assert len(pages) == 1


def test_plain_line_break_does_not_split():
    # A soft line break (<w:br/> with no w:type, or w:type="textWrapping")
    # must not be mistaken for a manual page break.
    doc = Document()
    p = doc.add_paragraph("Line one")
    p.add_run().add_break()  # plain <w:br/>, no type -> textWrapping
    p.add_run("Line two")

    buf = BytesIO()
    doc.save(buf)
    pages = convert_binary(BytesIO(buf.getvalue()))

    assert len(pages) == 1
