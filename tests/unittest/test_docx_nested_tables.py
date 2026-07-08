# Copyright (c) Opendatalab. All rights reserved.
import sys
import types
import zipfile
from io import BytesIO

from lxml import etree


office_image = types.ModuleType("mineru.backend.utils.office_image")
office_image.serialize_office_image = lambda *args, **kwargs: None
sys.modules.setdefault("mineru.backend.utils.office_image", office_image)

office_chart = types.ModuleType("mineru.backend.utils.office_chart")
office_chart.extract_chart_html_from_ooxml = lambda *args, **kwargs: ""
sys.modules.setdefault("mineru.backend.utils.office_chart", office_chart)

from docx import Document

from mineru.model.docx.docx_converter import DocxConverter


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NS = {"w": W_NS, "m": M_NS}


def _save_docx_bytes(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _convert_docx_bytes(file_bytes: bytes) -> list[dict]:
    converter = DocxConverter()
    converter.convert(BytesIO(file_bytes))
    return [block for page in converter.pages for block in page]


def _insert_omml_paragraph_into_first_body_table_cell(file_bytes: bytes) -> bytes:
    math_para = etree.fromstring(
        f"""
        <w:p xmlns:w="{W_NS}" xmlns:m="{M_NS}">
            <w:r>
                <m:oMath>
                    <m:r><m:t>FORMULA_X</m:t></m:r>
                </m:oMath>
            </w:r>
        </w:p>
        """.strip()
    )

    in_buffer = BytesIO(file_bytes)
    out_buffer = BytesIO()
    with zipfile.ZipFile(in_buffer, "r") as zin, zipfile.ZipFile(
        out_buffer, "w", zipfile.ZIP_DEFLATED
    ) as zout:
        for item in zin.infolist():
            data = zin.read(item.filename)
            if item.filename == "word/document.xml":
                root = etree.fromstring(data)
                outer_table = root.xpath("./w:body/w:tbl", namespaces=NS)[0]
                outer_cell = outer_table.xpath("./w:tr/w:tc", namespaces=NS)[0]
                for index, child in enumerate(list(outer_cell)):
                    if etree.QName(child).localname == "tbl":
                        outer_cell.insert(index, math_para)
                        break
                data = etree.tostring(
                    root,
                    encoding="utf-8",
                    xml_declaration=True,
                    standalone=True,
                )
            zout.writestr(item, data)
    return out_buffer.getvalue()


def test_nested_table_with_list_and_formula_is_not_dropped():
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    cell = outer.cell(0, 0)
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.style = "List Bullet"
    paragraph.add_run("bullet item")
    inner = cell.add_table(rows=1, cols=1)
    inner.cell(0, 0).text = "inner cell"

    file_bytes = _insert_omml_paragraph_into_first_body_table_cell(_save_docx_bytes(doc))
    blocks = _convert_docx_bytes(file_bytes)
    tables = [block for block in blocks if block["type"] == "table"]

    assert len(tables) == 1
    html = tables[0]["content"]
    assert html.count("<table") == 2
    assert "bullet item" in html
    assert "inner cell" in html
    assert "FORMULA" in html


def test_plain_nested_table_remains_embedded_in_parent_table_html():
    doc = Document()
    outer = doc.add_table(rows=1, cols=1)
    outer.cell(0, 0).text = "outer before"
    inner = outer.cell(0, 0).add_table(rows=1, cols=2)
    inner.cell(0, 0).text = "inner a"
    inner.cell(0, 1).text = "inner b"
    outer.cell(0, 0).add_paragraph("outer after")

    blocks = _convert_docx_bytes(_save_docx_bytes(doc))
    tables = [block for block in blocks if block["type"] == "table"]

    assert len(tables) == 1
    html = tables[0]["content"]
    assert html.count("<table") == 2
    assert "outer before" in html
    assert "inner a" in html
    assert "inner b" in html
    assert "outer after" in html


def test_colspan_normalization_does_not_count_nested_table_rows_or_cells():
    html = (
        '<table>'
        '<tr><td colspan="2">outer header</td></tr>'
        '<tr><td>outer a</td><td>'
        '<table><tr><td>inner a</td><td>inner b</td></tr></table>'
        '</td></tr>'
        '</table>'
    )

    normalized = DocxConverter()._normalize_table_colspans(html)

    assert '<td colspan="2">outer header</td>' in normalized
