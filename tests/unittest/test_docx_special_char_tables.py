# Copyright (c) Opendatalab. All rights reserved.
import sys
import types
from io import BytesIO

from lxml import etree

office_image = types.ModuleType("mineru.backend.utils.office_image")
office_image.serialize_office_image = lambda *args, **kwargs: None
sys.modules.setdefault("mineru.backend.utils.office_image", office_image)

office_chart = types.ModuleType("mineru.backend.utils.office_chart")
office_chart.extract_chart_html_from_ooxml = lambda *args, **kwargs: ""
sys.modules.setdefault("mineru.backend.utils.office_chart", office_chart)

from docx import Document
from docx.oxml.ns import qn

from mineru.model.docx.docx_converter import DocxConverter


W_NS = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
M_NS = "http://schemas.openxmlformats.org/officeDocument/2006/math"
NO_BREAK_HYPHEN = "‑"  # NON-BREAKING HYPHEN


def _save_docx_bytes(doc: Document) -> bytes:
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


def _convert_docx_bytes(file_bytes: bytes) -> list[dict]:
    converter = DocxConverter()
    converter.convert(BytesIO(file_bytes))
    return [block for page in converter.pages for block in page]


def _table_blocks(blocks: list[dict]) -> list[dict]:
    return [b for b in blocks if b["type"] == "table"]


def _build_docx_with_no_break_hyphen_table() -> bytes:
    """表格单元格同时含 <w:noBreakHyphen/> 与编号列表，是整表丢失的复现用例。"""
    doc = Document()
    doc.add_heading("1 项目概述", level=1)

    table = doc.add_table(rows=2, cols=2)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("NCI")
    etree.SubElement(r._r, qn("w:noBreakHyphen"))
    p.add_run("CTCAE 标准")
    p2 = table.cell(0, 1).paragraphs[0]
    p2.style = doc.styles["List Number"]
    p2.add_run("第一条编号内容")

    doc.add_heading("2 进度安排", level=1)
    t2 = doc.add_table(rows=1, cols=2)
    t2.cell(0, 0).paragraphs[0].add_run("里程碑")
    t2.cell(0, 1).paragraphs[0].add_run("2026-08-01")
    return _save_docx_bytes(doc)


def _build_docx_with_sym_table() -> bytes:
    """单元格含已映射 <w:sym>（(Symbol, 0x0022) → ∀）与编号列表的表格。"""
    doc = Document()
    doc.add_heading("符号表格", level=1)
    table = doc.add_table(rows=1, cols=2)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("符号")
    etree.SubElement(
        r._r,
        qn("w:sym"),
        {qn("w:char"): "0022", qn("w:font"): "Symbol"},
    )
    p.add_run("结束")
    p2 = table.cell(0, 1).paragraphs[0]
    p2.style = doc.styles["List Number"]
    p2.add_run("编号项")
    return _save_docx_bytes(doc)


def _build_docx_with_unmapped_sym_table() -> bytes:
    """单元格含未映射 <w:sym>（(Symbol, 0xF0B7)，Mammoth 渲染为空）的表格。"""
    doc = Document()
    doc.add_heading("符号表格", level=1)
    table = doc.add_table(rows=1, cols=2)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("符号")
    etree.SubElement(
        r._r,
        qn("w:sym"),
        {qn("w:char"): "F0B7", qn("w:font"): "Symbol"},
    )
    p.add_run("结束")
    p2 = table.cell(0, 1).paragraphs[0]
    p2.style = doc.styles["List Number"]
    p2.add_run("编号项")
    return _save_docx_bytes(doc)


def _build_docx_with_plain_table() -> bytes:
    """普通表格对照用例。"""
    doc = Document()
    doc.add_heading("对照", level=1)
    table = doc.add_table(rows=1, cols=2)
    table.cell(0, 0).paragraphs[0].add_run("APTT")
    table.cell(0, 1).paragraphs[0].add_run("部分凝血酶时间")
    return _save_docx_bytes(doc)


def test_no_break_hyphen_table_not_lost():
    blocks = _convert_docx_bytes(_build_docx_with_no_break_hyphen_table())
    tables = _table_blocks(blocks)
    assert len(tables) == 2

    first_html = tables[0]["content"]
    # 不间断连字符经 Mammoth 渲染保留，证明表格走了完整预解析路径而非被丢弃
    assert f"NCI{NO_BREAK_HYPHEN}CTCAE" in first_html
    # 单元格内编号列表由 Mammoth 渲染为列表，而非回退后的纯文本
    assert "<li>" in first_html
    assert "第一条编号内容" in first_html

    second_html = tables[1]["content"]
    assert "里程碑" in second_html
    assert "2026-08-01" in second_html


def test_mapped_sym_table_takes_mammoth_path():
    blocks = _convert_docx_bytes(_build_docx_with_sym_table())
    tables = _table_blocks(blocks)
    assert len(tables) == 1
    html = tables[0]["content"]
    # Mammoth 按 dingbats 映射渲染 (Symbol, 0x0022) → ∀（U+2200）
    assert "∀" in html
    # 编号列表由 Mammoth 渲染，证明表格走了完整预解析路径
    assert "<li>" in html


def test_unmapped_sym_table_with_numbering_not_lost():
    """未映射的 w:sym Mammoth 渲染为空，签名也必须为空，否则表格会被丢弃。"""
    blocks = _convert_docx_bytes(_build_docx_with_unmapped_sym_table())
    tables = _table_blocks(blocks)
    assert len(tables) == 1
    html = tables[0]["content"]
    assert "符号" in html and "结束" in html
    assert "<li>" in html


def test_plain_table_still_emitted():
    blocks = _convert_docx_bytes(_build_docx_with_plain_table())
    tables = _table_blocks(blocks)
    assert len(tables) == 1
    assert "APTT" in tables[0]["content"]


def test_xml_table_signature_renders_special_chars():
    """签名文本必须与 Mammoth 渲染一致：noBreakHyphen→U+2011、softHyphen→U+00AD。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("NCI")
    etree.SubElement(r._r, qn("w:noBreakHyphen"))
    p.add_run("CTCAE")
    p2 = table.cell(0, 1).paragraphs[0]
    r2 = p2.add_run("软连")
    etree.SubElement(r2._r, qn("w:softHyphen"))
    p2.add_run("字符")

    sig = DocxConverter._xml_table_signature(table._tbl)
    assert f"NCI{NO_BREAK_HYPHEN}CTCAE" in sig["text"]
    assert "软连­字符" in sig["text"]


def test_xml_table_signature_renders_sym_like_mammoth():
    """w:sym 按 Mammoth dingbats 映射渲染：已映射输出字符，未映射输出空。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("映射")
    etree.SubElement(
        r._r,
        qn("w:sym"),
        {qn("w:char"): "0022", qn("w:font"): "Symbol"},
    )
    p.add_run("结束")
    p2 = table.cell(0, 1).paragraphs[0]
    r2 = p2.add_run("未映射")
    etree.SubElement(
        r2._r,
        qn("w:sym"),
        {qn("w:char"): "F0B7", qn("w:font"): "Symbol"},
    )

    sig = DocxConverter._xml_table_signature(table._tbl)
    # (Symbol, 0x0022) → ∀（U+2200）
    assert "映射∀结束" in sig["text"]
    # (Symbol, 0xF0B7) 未映射，Mammoth 渲染为空，签名也必须为空
    assert "未映射" in sig["text"]
    assert "" not in sig["text"]


def test_xml_table_signature_excludes_omml_equation_text():
    """OMML 公式的 m:t 不计入签名，与 Mammoth 丢弃公式的行为保持一致。"""
    doc = Document()
    table = doc.add_table(rows=1, cols=2)
    p = table.cell(0, 0).paragraphs[0]
    r = p.add_run("正文")
    omath = etree.SubElement(r._r, qn("m:oMath"))
    omr = etree.SubElement(omath, qn("m:r"))
    etree.SubElement(omr, qn("m:t")).text = "E=mc2"
    table.cell(0, 1).paragraphs[0].add_run("单元格")

    sig = DocxConverter._xml_table_signature(table._tbl)
    assert "正文" in sig["text"]
    assert "E=mc2" not in sig["text"]
