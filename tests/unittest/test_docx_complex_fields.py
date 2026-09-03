from __future__ import annotations

from io import BytesIO

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.text.paragraph import Paragraph

from mineru.backend.analyze import doc_analyze
from mineru.backend.postprocess.inline import inline_plain_text
from mineru.model.flash import DocxModel
from mineru.model.flash._shared.spans import inline_span_plain_text
from mineru.types import IndexBlock, ParagraphTitleBlock, TextBlock
from _span_test_utils import inline_urls


def _append_field_char(paragraph: Paragraph, field_type: str) -> None:
    """向段落追加一个复杂字段边界 run。"""
    run = OxmlElement("w:r")
    field_char = OxmlElement("w:fldChar")
    field_char.set(qn("w:fldCharType"), field_type)
    run.append(field_char)
    paragraph._p.append(run)


def _append_instruction(paragraph: Paragraph, instruction: str) -> None:
    """向段落追加复杂字段指令 run。"""
    run = OxmlElement("w:r")
    instruction_element = OxmlElement("w:instrText")
    instruction_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    instruction_element.text = instruction
    run.append(instruction_element)
    paragraph._p.append(run)


def _append_result_text(
    paragraph: Paragraph,
    text: str,
    *,
    bold: bool = False,
) -> None:
    """向字段结果追加一个可见文本 run。"""
    run = OxmlElement("w:r")
    if bold:
        run_properties = OxmlElement("w:rPr")
        run_properties.append(OxmlElement("w:b"))
        run.append(run_properties)
    text_element = OxmlElement("w:t")
    if text != text.strip():
        text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
    text_element.text = text
    run.append(text_element)
    paragraph._p.append(run)


def _append_native_hyperlink(
    paragraph: Paragraph,
    target: str,
    parts: list[tuple[str, bool]],
) -> None:
    """向段落追加由多个格式 run 组成的真实外部超链接。"""

    relationship_id = paragraph.part.relate_to(
        target,
        RELATIONSHIP_TYPE.HYPERLINK,
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship_id)
    for text, bold in parts:
        run = OxmlElement("w:r")
        if bold:
            run_properties = OxmlElement("w:rPr")
            run_properties.append(OxmlElement("w:b"))
            run.append(run_properties)
        text_element = OxmlElement("w:t")
        if text != text.strip():
            text_element.set("{http://www.w3.org/XML/1998/namespace}space", "preserve")
        text_element.text = text
        run.append(text_element)
        hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _append_result_tab(paragraph: Paragraph) -> None:
    """向字段结果追加 Word 制表符 run。"""
    run = OxmlElement("w:r")
    run.append(OxmlElement("w:tab"))
    paragraph._p.append(run)


def _append_toc_complex_field(
    paragraph: Paragraph,
    *,
    title: str,
    page_number: str,
    anchor: str,
    include_outer_toc: bool,
    split_hyperlink_instruction: bool = False,
) -> None:
    """构造 WPS 常见的 TOC、HYPERLINK 与 PAGEREF 嵌套复杂域。"""
    if include_outer_toc:
        _append_field_char(paragraph, "begin")
        _append_instruction(paragraph, ' TOC \\o "1-1" \\h \\z \\u ')
        _append_field_char(paragraph, "separate")

    _append_field_char(paragraph, "begin")
    if split_hyperlink_instruction:
        split_at = max(1, len(anchor) // 2)
        _append_instruction(paragraph, " HYPER")
        _append_instruction(paragraph, f'LINK \\l "{anchor[:split_at]}')
        _append_instruction(paragraph, f'{anchor[split_at:]}" ')
    else:
        _append_instruction(paragraph, f' HYPERLINK \\l "{anchor}" ')
    _append_field_char(paragraph, "separate")
    _append_result_text(paragraph, title)
    _append_result_tab(paragraph)
    _append_field_char(paragraph, "begin")
    _append_instruction(paragraph, f" PAGEREF {anchor} \\h ")
    _append_field_char(paragraph, "separate")
    _append_result_text(paragraph, page_number)
    _append_field_char(paragraph, "end")
    _append_field_char(paragraph, "end")

    if include_outer_toc:
        _append_field_char(paragraph, "end")


def _attach_bookmark(paragraph: Paragraph, anchor: str, bookmark_id: int) -> None:
    """把 bookmark 包围段落现有正文，模拟 TOC 的真实跳转目标。"""
    start = OxmlElement("w:bookmarkStart")
    start.set(qn("w:id"), str(bookmark_id))
    start.set(qn("w:name"), anchor)
    end = OxmlElement("w:bookmarkEnd")
    end.set(qn("w:id"), str(bookmark_id))
    insert_at = 1 if paragraph._p.pPr is not None else 0
    paragraph._p.insert(insert_at, start)
    paragraph._p.append(end)


def _build_complex_toc_docx() -> bytes:
    """生成一个包含匹配目标和缺失目标的最小复杂域目录 DOCX。"""
    document = Document()
    toc_style = document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)

    first = document.add_paragraph(style=toc_style)
    _append_toc_complex_field(
        first,
        title="一、建设项目基本情况",
        page_number="1",
        anchor="_TocTarget",
        include_outer_toc=True,
        split_hyperlink_instruction=True,
    )
    second = document.add_paragraph(style=toc_style)
    _append_toc_complex_field(
        second,
        title="二、缺失章节",
        page_number="9",
        anchor="_TocMissing",
        include_outer_toc=False,
    )

    heading = document.add_heading("一、建设项目基本情况", level=1)
    _attach_bookmark(heading, "_TocTarget", 7)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _build_external_bookmark_field_docx() -> bytes:
    """生成同时包含外部文档地址与 bookmark switch 的复杂字段。"""

    document = Document()
    paragraph = document.add_paragraph()
    _append_field_char(paragraph, "begin")
    _append_instruction(
        paragraph,
        ' HYPERLINK "chapter.docx" \\l "Section1" ',
    )
    _append_field_char(paragraph, "separate")
    _append_result_text(paragraph, "Open ")
    _append_result_text(paragraph, "section", bold=True)
    _append_field_char(paragraph, "end")
    output = BytesIO()
    document.save(output)
    return output.getvalue()


def _build_multi_alias_toc_docx() -> bytes:
    """生成两个被引用 TOC bookmark 指向同一正文段落的 DOCX。"""

    document = Document()
    toc_style = document.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    for title, anchor in (
        ("Primary entry", "_TocPrimary"),
        ("Alias entry", "_TocAlias"),
    ):
        paragraph = document.add_paragraph(style=toc_style)
        _append_toc_complex_field(
            paragraph,
            title=title,
            page_number="1",
            anchor=anchor,
            include_outer_toc=False,
            split_hyperlink_instruction=True,
        )

    heading = document.add_heading("Shared target", level=1)
    insert_at = 1 if heading._p.pPr is not None else 0
    for offset, (bookmark_id, anchor) in enumerate(
        ((10, "_TocPrimary"), (11, "_TocAlias")),
    ):
        start = OxmlElement("w:bookmarkStart")
        start.set(qn("w:id"), str(bookmark_id))
        start.set(qn("w:name"), anchor)
        heading._p.insert(insert_at + offset, start)
        end = OxmlElement("w:bookmarkEnd")
        end.set(qn("w:id"), str(bookmark_id))
        heading._p.append(end)

    output = BytesIO()
    document.save(output)
    return output.getvalue()


def test_nested_complex_toc_fields_preserve_titles_and_strict_index_targets() -> None:
    """验证拆分且嵌套的复杂域不会只剩页码，并按真实正文目标收敛目录 anchor。"""
    file_bytes = _build_complex_toc_docx()
    model_pages = DocxModel().predict(BytesIO(file_bytes))
    raw_index = next(block for page in model_pages for block in page if block["type"] == "index")

    assert [inline_span_plain_text(child["content"]) for child in raw_index["content"]] == [
        "一、建设项目基本情况\t1",
        "二、缺失章节\t9",
    ]
    assert [child.get("anchor") for child in raw_index["content"]] == ["_TocTarget", "_TocMissing"]

    middle_json, _ = doc_analyze(file_bytes, effort="flash", parse_mode="auto", file_suffix="docx")
    typed_index = next(block for page in middle_json.pages for block in page.blocks if isinstance(block, IndexBlock))
    assert isinstance(typed_index.content[0], ParagraphTitleBlock)
    assert typed_index.content[0].anchor == "_TocTarget"
    assert inline_plain_text(typed_index.content[0].content) == "一、建设项目基本情况\t1"
    assert isinstance(typed_index.content[1], TextBlock)
    assert typed_index.content[1].anchor is None
    assert inline_plain_text(typed_index.content[1].content) == "二、缺失章节\t9"


def test_external_complex_hyperlink_preserves_bookmark_switch() -> None:
    """验证外部复杂字段把文档地址与 bookmark 合成为同一个链接目标。"""

    file_bytes = _build_external_bookmark_field_docx()
    model_pages = DocxModel().predict(BytesIO(file_bytes))
    raw_text = next(block for page in model_pages for block in page if block["type"] == "text")

    assert inline_span_plain_text(raw_text["content"]) == "Open section"
    assert inline_urls(raw_text["content"]) == ["chapter.docx#Section1"]
    assert [child["styles"] for child in raw_text["content"][0]["content"]] == [[], ["bold"]]

    middle_json, _ = doc_analyze(file_bytes, effort="flash", parse_mode="auto", file_suffix="docx")
    typed_text = next(block for page in middle_json.pages for block in page.blocks if isinstance(block, TextBlock))
    assert inline_plain_text(typed_text.content) == "Open section"
    assert inline_urls(typed_text.content) == ["chapter.docx#Section1"]


def test_split_toc_bookmark_aliases_collapse_to_one_target_anchor() -> None:
    """验证拆分字段引用的同段落 bookmark aliases 收敛为唯一公开 anchor。"""

    file_bytes = _build_multi_alias_toc_docx()
    model_pages = DocxModel().predict(BytesIO(file_bytes))
    raw_index = next(block for page in model_pages for block in page if block["type"] == "index")
    raw_heading = next(
        block
        for page in model_pages
        for block in page
        if block["type"] == "paragraph_title" and inline_span_plain_text(block["content"]) == "Shared target"
    )

    assert raw_heading["anchor"] == "_TocPrimary"
    assert [child.get("anchor") for child in raw_index["content"]] == [
        "_TocPrimary",
        "_TocPrimary",
    ]

    middle_json, _ = doc_analyze(
        file_bytes,
        effort="flash",
        parse_mode="auto",
        file_suffix="docx",
    )
    typed_index = next(block for page in middle_json.pages for block in page.blocks if isinstance(block, IndexBlock))
    typed_heading = next(
        block
        for page in middle_json.pages
        for block in page.blocks
        if isinstance(block, ParagraphTitleBlock) and inline_plain_text(block.content) == "Shared target"
    )
    assert typed_heading.anchor == "_TocPrimary"
    assert [child.anchor for child in typed_index.content] == [
        "_TocPrimary",
        "_TocPrimary",
    ]


def test_native_hyperlink_preserves_spaces_between_formatted_runs() -> None:
    """验证真实超链接跨格式 run 时保留内部空格并维持单一链接目标。"""

    document = Document()
    paragraph = document.add_paragraph()
    _append_native_hyperlink(
        paragraph,
        "https://example.test/section",
        [("Open ", False), ("section", True)],
    )
    output = BytesIO()
    document.save(output)

    model_pages = DocxModel().predict(BytesIO(output.getvalue()))
    raw_text = next(block for page in model_pages for block in page if block["type"] == "text")

    assert inline_span_plain_text(raw_text["content"]) == "Open section"
    assert inline_urls(raw_text["content"]) == ["https://example.test/section"]
    assert [child["styles"] for child in raw_text["content"][0]["content"]] == [[], ["bold"]]
