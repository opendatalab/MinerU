from io import BytesIO
import re

from docx import Document
from PIL import Image

from mineru.backend.office.docx_analyze import office_docx_analyze
from mineru.model.docx.main import convert_binary
from mineru.utils.enum_class import BlockType


class RecordingImageWriter:
    def __init__(self):
        self.files = {}

    def write(self, path: str, data: bytes) -> None:
        self.files[path] = data


def _write_png(path, color: tuple[int, int, int]) -> bytes:
    image = Image.new("RGB", (24, 24), color)
    image.save(path, format="PNG")
    return path.read_bytes()


def _build_docx_with_body_and_table_images(tmp_path) -> tuple[bytes, bytes]:
    body_image_path = tmp_path / "body.png"
    table_image_path = tmp_path / "table.png"
    _write_png(body_image_path, (255, 0, 0))
    table_image_bytes = _write_png(table_image_path, (0, 255, 0))

    document = Document()
    document.add_paragraph("body image:")
    document.add_picture(str(body_image_path))

    table = document.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    cell.text = "table image:"
    cell.add_paragraph().add_run().add_picture(str(table_image_path))

    output = BytesIO()
    document.save(output)
    return output.getvalue(), table_image_bytes


def test_docx_table_preparse_writes_table_images_without_base64(tmp_path):
    docx_bytes, table_image_bytes = _build_docx_with_body_and_table_images(tmp_path)
    image_writer = RecordingImageWriter()

    pages = convert_binary(BytesIO(docx_bytes), image_writer=image_writer)

    table_html = next(
        block["content"]
        for page in pages
        for block in page
        if block["type"] == BlockType.TABLE
    )
    image_block = next(
        block
        for page in pages
        for block in page
        if block["type"] == BlockType.IMAGE
    )
    src_match = re.search(r'src="([^"]+)"', table_html)

    assert "base64," not in table_html
    assert src_match is not None
    assert src_match.group(1) in image_writer.files
    assert image_writer.files[src_match.group(1)] == table_image_bytes
    assert "content" not in image_block
    assert image_block["image_path"] in image_writer.files
    assert len(image_writer.files) == 2


def test_docx_office_middle_json_keeps_image_paths_without_base64(tmp_path):
    docx_bytes, _ = _build_docx_with_body_and_table_images(tmp_path)
    image_writer = RecordingImageWriter()

    middle_json, _ = office_docx_analyze(docx_bytes, image_writer=image_writer)

    image_spans = []
    table_spans = []
    for page_info in middle_json["pdf_info"]:
        for block in page_info["para_blocks"]:
            for sub_block in block.get("blocks", []):
                for line in sub_block.get("lines", []):
                    for span in line.get("spans", []):
                        if span["type"] == "image":
                            image_spans.append(span)
                        if span["type"] == "table":
                            table_spans.append(span)

    assert image_spans
    assert table_spans
    assert all("image_base64" not in span for span in image_spans)
    assert all(span.get("image_path") in image_writer.files for span in image_spans)
    assert all("base64," not in span["html"] for span in table_spans)
