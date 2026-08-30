from __future__ import annotations
from _span_test_utils import inline as _inline

from copy import deepcopy
from datetime import datetime, timezone
from io import BytesIO
from inspect import signature

from docx import Document
from PIL import Image
import pytest

import mineru.render as render_module
from mineru.render import (
    ContentListRenderOptions,
    ContentListV2RenderOptions,
    DocxRenderError,
    DocxRenderOptions,
    EpubRenderOptions,
    HtmlRenderOptions,
    MarkdownRenderOptions,
    PdfRenderOptions,
    RenderFormat,
    RenderMode,
    StructuredContentRenderOptions,
    render,
)
from mineru.types import ImageBlock, ImageBodyBlock, MiddleJson, PageBlock, PageInfo, TextBlock


def _middle(*pages: PageInfo) -> MiddleJson:
    """构造无需 PDF bbox 的最小严格 MiddleJson。"""
    return MiddleJson(
        pages=list(pages),
        is_full_document=True,
        file_suffix="docx",
        effort="flash",
        parse_mode="txt",
        mineru_version="test",
    )


def _page(page_idx: int, *blocks: PageBlock) -> PageInfo:
    """按调用方顺序构造一页严格 MiddleJson 内容。"""
    return PageInfo(page_idx=page_idx, blocks=list(blocks))


def _png_bytes() -> bytes:
    """生成可被 Pillow 与 python-docx 完整读取的 PNG。"""
    output = BytesIO()
    Image.new("RGB", (2, 2), (30, 60, 90)).save(output, format="PNG")
    return output.getvalue()


def test_unified_render_dispatches_all_native_output_types_without_mutation() -> None:
    """验证统一入口分发全部原生结果且不修改输入。"""
    middle = _middle(_page(0, TextBlock(type="text", index=0, content=_inline("hello"))))
    original = deepcopy(middle)

    markdown = render(middle, RenderFormat.MARKDOWN)
    html = render(middle, RenderFormat.HTML)
    docx = render(middle, RenderFormat.DOCX)
    epub = render(
        middle,
        RenderFormat.EPUB,
        options=EpubRenderOptions(modified_at=datetime(2026, 1, 2, tzinfo=timezone.utc)),
    )
    pdf = render(middle, RenderFormat.PDF)
    structured_content = render(middle, RenderFormat.STRUCTURED_CONTENT)
    content_list = render(middle, RenderFormat.CONTENT_LIST)
    content_list_v2 = render(middle, RenderFormat.CONTENT_LIST_V2)

    assert markdown == "hello"
    assert isinstance(html, str) and "<html" in html and "hello" in html
    assert isinstance(docx, bytes) and docx.startswith(b"PK\x03\x04")
    assert isinstance(epub, bytes) and epub.startswith(b"PK\x03\x04")
    assert isinstance(pdf, bytes) and pdf.startswith(b"%PDF-")
    assert Document(BytesIO(docx)).paragraphs[0].text == "hello"
    assert structured_content["pages"][0]["blocks"][0]["content"] == "hello"
    assert content_list == [{"type": "text", "text": "hello", "page_idx": 0}]
    assert content_list_v2 == [
        [{"type": "paragraph", "content": {"paragraph_content": [{"type": "text", "content": "hello"}]}}]
    ]
    assert middle == original


def test_unified_render_forwards_format_specific_options() -> None:
    """验证 Markdown/HTML 模式与各格式专属资源选项分别透传。"""
    middle = _middle(
        _page(0, TextBlock(type="text", index=0, content=_inline("first-"))),
        _page(1, TextBlock(type="text", index=0, content=_inline("second"), continues_prev=True)),
    )
    image_middle = _middle(
        _page(
            0,
            ImageBlock(
                type="image",
                index=0,
                content=[
                    ImageBodyBlock(
                        type="image_body",
                        index=0,
                        content="",
                        image_path="images/a b.png",
                    )
                ],
            ),
        )
    )

    markdown = render(
        middle,
        RenderFormat.MARKDOWN,
        options=MarkdownRenderOptions(mode=RenderMode.FULL),
    )
    html = render(
        middle,
        RenderFormat.HTML,
        options=HtmlRenderOptions(mode=RenderMode.FULL, standalone=False),
    )
    structured_content = render(
        image_middle,
        RenderFormat.STRUCTURED_CONTENT,
        options=StructuredContentRenderOptions(asset_base_url="https://cdn.example/doc"),
    )
    content_list = render(
        image_middle,
        RenderFormat.CONTENT_LIST,
        options=ContentListRenderOptions(asset_base_url="https://cdn.example/doc"),
    )
    content_list_v2 = render(
        image_middle,
        RenderFormat.CONTENT_LIST_V2,
        options=ContentListV2RenderOptions(asset_base_url="https://cdn.example/doc"),
    )
    image_markdown = render(
        image_middle,
        RenderFormat.MARKDOWN,
        options=MarkdownRenderOptions(asset_base_url="https://cdn.example/doc"),
    )
    html_document = render(
        image_middle,
        RenderFormat.HTML,
        options=HtmlRenderOptions(
            asset_base_url="https://cdn.example/doc",
            document_title="Unified Render",
        ),
    )
    pdf = render(
        middle,
        RenderFormat.PDF,
        options=PdfRenderOptions(document_title="Unified Render"),
    )
    original_tree = render(middle, RenderFormat.STRUCTURED_CONTENT)
    original_list = render(middle, RenderFormat.CONTENT_LIST)
    original_list_v2 = render(middle, RenderFormat.CONTENT_LIST_V2)

    assert "\n\n---\n\n" in markdown
    assert html.startswith('<article class="mineru-document mineru-document--full" ')
    assert 'data-mineru-html-version="1" data-render-mode="full"' in html
    assert "<!doctype html>" not in html
    assert structured_content["pages"][0]["blocks"][0]["image_source"] == ("https://cdn.example/doc/images/a%20b.png")
    assert content_list[0]["img_path"] == "https://cdn.example/doc/images/a%20b.png"
    assert content_list_v2[0][0]["content"]["image_source"] == {"path": "https://cdn.example/doc/images/a%20b.png"}
    assert "https://cdn.example/doc/images/a%20b.png" in image_markdown
    assert "<title>Unified Render</title>" in html_document
    assert 'src="https://cdn.example/doc/images/a%20b.png"' in html_document
    assert pdf == render_module.render_pdf(middle, document_title="Unified Render")
    assert [page["blocks"][0]["content"] for page in original_tree["pages"]] == ["first-", "second"]
    assert original_tree["pages"][1]["blocks"][0]["continues_prev"] is True
    assert [(item["text"], item["page_idx"]) for item in original_list] == [("first-", 0), ("second", 1)]
    assert [page[0]["content"]["paragraph_content"][0]["content"] for page in original_list_v2] == [
        "first-",
        "second",
    ]


def test_unified_docx_uses_typed_asset_resolver_and_propagates_public_errors() -> None:
    """验证 DOCX Options 传递素材解析器，并保留带定位的公共异常。"""
    requested_paths: list[str] = []

    def resolve_asset(relative_path: str) -> bytes:
        """记录统一入口请求的相对路径并返回有效 PNG。"""
        requested_paths.append(relative_path)
        return _png_bytes()

    middle = _middle(
        _page(
            7,
            ImageBlock(
                type="image",
                index=3,
                content=[
                    ImageBodyBlock(
                        type="image_body",
                        index=3,
                        content="description",
                        image_path="images/picture.png",
                    )
                ],
            ),
        )
    )

    docx = render(
        middle,
        RenderFormat.DOCX,
        options=DocxRenderOptions(asset_resolver=resolve_asset),
    )

    assert requested_paths == ["images/picture.png"]
    assert len(Document(BytesIO(docx)).inline_shapes) == 1
    with pytest.raises(DocxRenderError) as exc_info:
        render(middle, RenderFormat.DOCX)
    assert (exc_info.value.page_idx, exc_info.value.block_index, exc_info.value.block_type) == (7, 3, "image")


def test_unified_render_rejects_legacy_format_and_mismatched_options() -> None:
    """验证统一入口拒绝旧字典、字符串格式和跨格式 Options。"""
    middle = _middle(_page(0))

    with pytest.raises(TypeError, match="MiddleJson"):
        render(middle.to_dict(), RenderFormat.MARKDOWN)  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="RenderFormat"):
        render(middle, "markdown")  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="MarkdownRenderOptions"):
        render(middle, RenderFormat.MARKDOWN, options=HtmlRenderOptions())  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="StructuredContentRenderOptions"):
        render(middle, RenderFormat.STRUCTURED_CONTENT, options=DocxRenderOptions())  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="EpubRenderOptions"):
        render(middle, RenderFormat.EPUB, options=DocxRenderOptions())  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="PdfRenderOptions"):
        render(middle, RenderFormat.PDF, options=DocxRenderOptions())  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="ContentListRenderOptions"):
        render(middle, RenderFormat.CONTENT_LIST, options=DocxRenderOptions())  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="ContentListV2RenderOptions"):
        render(middle, RenderFormat.CONTENT_LIST_V2, options=ContentListRenderOptions())  # type: ignore[arg-type]


def test_public_options_validate_fields() -> None:
    """验证 Options 构造期的严格字段类型检查。"""
    assert [item.value for item in RenderFormat] == [
        "markdown",
        "html",
        "docx",
        "epub",
        "structured_content",
        "content_list",
        "content_list_v2",
        "pdf",
    ]

    with pytest.raises(TypeError, match="RenderMode"):
        MarkdownRenderOptions(mode="default")  # type: ignore[arg-type]
    assert "mode" in signature(MarkdownRenderOptions).parameters
    assert "mode" in signature(HtmlRenderOptions).parameters
    for options_type in (
        DocxRenderOptions,
        EpubRenderOptions,
        PdfRenderOptions,
        StructuredContentRenderOptions,
        ContentListRenderOptions,
        ContentListV2RenderOptions,
    ):
        assert "mode" not in signature(options_type).parameters
        with pytest.raises(TypeError, match="unexpected keyword argument 'mode'"):
            options_type(mode=RenderMode.DEFAULT)  # type: ignore[call-arg]
    with pytest.raises(TypeError, match="asset_base_url"):
        StructuredContentRenderOptions(asset_base_url=1)  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="asset_base_url"):
        ContentListRenderOptions(asset_base_url=1)  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="asset_base_url"):
        ContentListV2RenderOptions(asset_base_url=1)  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="standalone"):
        HtmlRenderOptions(standalone=1)  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="document_title"):
        HtmlRenderOptions(document_title=1)  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="asset_resolver"):
        DocxRenderOptions(asset_resolver="images")  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="authors"):
        EpubRenderOptions(authors=["Alice"])  # type: ignore[arg-type]
    with pytest.raises(ValueError, match="BCP 47"):
        EpubRenderOptions(language="zh CN")
    with pytest.raises(ValueError, match="timezone-aware"):
        EpubRenderOptions(modified_at=datetime(2026, 1, 2))
    with pytest.raises(ValueError, match="four digits"):
        EpubRenderOptions(modified_at=datetime(999, 1, 2, tzinfo=timezone.utc))
    with pytest.raises(TypeError, match="asset_resolver"):
        EpubRenderOptions(asset_resolver="images")  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="asset_resolver"):
        PdfRenderOptions(asset_resolver="images")  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="document_title"):
        PdfRenderOptions(document_title=1)  # type: ignore[arg-type]


def test_public_render_exposes_all_three_structured_output_names() -> None:
    """验证严格 render 公共面同时暴露树形输出和两个兼容 Content List。"""
    assert callable(render_module.render_structured_content)
    assert callable(render_module.render_content_list)
    assert callable(render_module.render_content_list_v2)
    assert callable(render_module.render_epub)
    assert callable(render_module.render_pdf)
    assert not hasattr(render_module, "MarkdownRenderMode")
    assert ContentListRenderOptions is render_module.ContentListRenderOptions
    assert ContentListV2RenderOptions is render_module.ContentListV2RenderOptions
    assert "CONTENT_LIST" in RenderFormat.__members__
    assert "CONTENT_LIST_V2" in RenderFormat.__members__
