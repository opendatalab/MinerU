from __future__ import annotations
from _span_test_utils import inline as _inline

import base64
from copy import deepcopy
from io import BytesIO
from unittest.mock import Mock
import zipfile

from docx import Document
from docx.shared import Mm, Twips
from lxml import etree
from PIL import Image
import pytest

from mineru.render import DocxRenderError, render_docx
from mineru.types import (
    ChartBlock,
    ChartBodyBlock,
    CodeBlock,
    EquationBlock,
    ImageBlock,
    ImageBodyBlock,
    IndexBlock,
    ListBlock,
    MiddleJson,
    PageAuxTextBlock,
    PageBlock,
    PageFootnoteBlock,
    PageInfo,
    ParagraphTitleBlock,
    TableAnnotationBlock,
    TableBlock,
    TableBodyBlock,
    TextBlock,
)


def _middle(*pages: PageInfo) -> MiddleJson:
    """构造最小严格 MiddleJson 测试对象。"""
    return MiddleJson(
        pages=list(pages),
        is_full_document=True,
        file_suffix="docx",
        effort="flash",
        parse_mode="txt",
        mineru_version="test",
    )


def _page(page_idx: int, *blocks: PageBlock) -> PageInfo:
    """构造保持调用方 block 顺序的严格页面。"""
    return PageInfo(page_idx=page_idx, blocks=list(blocks))


def _png_bytes(*, size: tuple[int, int] = (12, 8)) -> bytes:
    """生成可被 Pillow 和 python-docx 完整读取的 PNG。"""
    output = BytesIO()
    Image.new("RGB", size, (30, 60, 90)).save(output, format="PNG")
    return output.getvalue()


def _png_uri(*, size: tuple[int, int] = (12, 8)) -> str:
    """生成严格 PNG data URI。"""
    payload = base64.b64encode(_png_bytes(size=size)).decode("ascii")
    return f"data:image/png;base64,{payload}"


def _generated_svg_uri(*, logical_size: tuple[int, int] = (12, 8), fallback_size: tuple[int, int] = (96, 64)) -> str:
    """生成带高密度 PNG fallback 的 MinerU SVG data URI。"""
    width, height = logical_size
    fallback = base64.b64encode(_png_bytes(size=fallback_size)).decode("ascii")
    svg = (
        f'<svg xmlns="http://www.w3.org/2000/svg" width="{width}" height="{height}" '
        f'viewBox="0 0 {width} {height}" data-mineru-generated="wmf-emf">'
        f'<metadata id="mineru-raster-fallback" data-mime="image/png">{fallback}</metadata>'
        f'<path d="M 0 0 L {width} 0 L {width} {height} Z" fill="#000000"/>'
        "</svg>"
    ).encode()
    return f"data:image/svg+xml;base64,{base64.b64encode(svg).decode('ascii')}"


def _part(docx_bytes: bytes, name: str) -> str:
    """读取 DOCX ZIP 中一个 XML part。"""
    with zipfile.ZipFile(BytesIO(docx_bytes)) as archive:
        return archive.read(name).decode("utf-8")


def test_public_contract_returns_reopenable_docx_without_mutation() -> None:
    """验证严格入口、可重开 bytes 和输入无副作用。"""
    middle = _middle(_page(0, TextBlock(type="text", index=0, content=_inline("hello"))))
    original = deepcopy(middle)

    result = render_docx(middle)

    assert result.startswith(b"PK\x03\x04")
    assert Document(BytesIO(result)).paragraphs[0].text == "hello"
    assert middle == original
    with pytest.raises(TypeError, match="MiddleJson"):
        render_docx(middle.to_dict())  # type: ignore[arg-type]
    with pytest.raises(TypeError, match="unexpected keyword argument 'mode'"):
        render_docx(middle, mode="full")  # type: ignore[arg-type]


def test_xml_incompatible_text_is_replaced_with_visible_marker() -> None:
    """验证真实语料中的 XML 控制字符会转为 U+FFFD，而不是中止或静默删除。"""
    middle = _middle(_page(0, TextBlock(type="text", index=0, content=_inline("before\x01after"))))

    document = Document(BytesIO(render_docx(middle)))

    assert document.paragraphs[0].text == "before\ufffdafter"


def test_heading_bookmark_forward_index_link_and_rich_inline_ooxml() -> None:
    """验证 Heading、前向目录链接、外链、组合样式和行内 OMML。"""
    index = IndexBlock(
        type="index",
        index=0,
        content=[
            ParagraphTitleBlock(
                type="paragraph_title",
                level=2,
                anchor="section a",
                content=_inline("Section\t12"),
            )
        ],
    )
    text = TextBlock(
        type="text",
        index=1,
        content=[
            {
                "type": "text",
                "content": "Styled",
                "styles": ["bold", "italic", "underline", "strikethrough", "emphasis"],
            },
            {"type": "text", "content": "2", "styles": ["superscript"]},
            {"type": "text", "content": "i", "styles": ["subscript"]},
            {"type": "hyperlink", "url": "https://example.com/a", "content": _inline("Link")},
            {"type": "equation_inline", "content": "x^2"},
        ],
    )
    title = ParagraphTitleBlock(
        type="paragraph_title",
        index=2,
        level=2,
        anchor="section a",
        content=_inline("Section"),
    )

    result = render_docx(_middle(_page(0, index, text, title)))
    document_xml = _part(result, "word/document.xml")
    relationships = _part(result, "word/_rels/document.xml.rels")

    assert 'w:pStyle w:val="Heading2"' in document_xml
    assert 'w:bookmarkStart w:id="0" w:name="section_a"' in document_xml
    assert 'w:hyperlink w:anchor="section_a"' in document_xml
    assert "https://example.com/a" in relationships
    assert "<w:b" in document_xml and "<w:i" in document_xml
    assert "<w:u" in document_xml and "<w:strike" in document_xml
    assert '<w:vertAlign w:val="superscript"' in document_xml
    assert '<w:vertAlign w:val="subscript"' in document_xml
    assert '<w:em w:val="underDot"' in document_xml
    assert "<m:oMath" in document_xml


def test_docx_uses_page_footnote_bookmark_and_style() -> None:
    """验证固定默认 Word 输出页面脚注书签、内部链接和专用样式。"""
    middle = _middle(
        _page(
            0,
            TextBlock(
                type="text",
                index=0,
                content=[
                    {"type": "text", "content": "See "},
                    {"type": "hyperlink", "url": "#note-one", "content": _inline("[1]")},
                    {"type": "text", "content": " and "},
                    {"type": "hyperlink", "url": "#missing-note", "content": _inline("[x]")},
                    {"type": "text", "content": "."},
                ],
            ),
            PageFootnoteBlock(
                type="page_footnote",
                index=1,
                content=_inline("Footnote body."),
                anchor="note-one",
            ),
        )
    )

    result = render_docx(middle)
    document_xml = _part(result, "word/document.xml")
    relationships = _part(result, "word/_rels/document.xml.rels")
    document = Document(BytesIO(result))
    footnote = next(paragraph for paragraph in document.paragraphs if paragraph.text == "Footnote body.")

    assert 'w:bookmarkStart w:id="0" w:name="note_one"' in document_xml
    assert 'w:hyperlink w:anchor="note_one"' in document_xml
    assert document_xml.count("<w:hyperlink") == 1
    assert footnote.style.name == "MinerU Footnote"
    assert "missing-note" not in relationships
    assert "#note-one" not in relationships


def test_visible_styled_boundary_spaces_use_nbsp_without_mutating_input() -> None:
    """验证可见样式的边界空格转为等量 NBSP，普通样式和内部空格保持原样。"""
    content = [
        {"type": "text", "content": "  left", "styles": ["underline"]},
        {"type": "text", "content": "|"},
        {"type": "text", "content": "right  ", "styles": ["strikethrough"]},
        {"type": "text", "content": "|"},
        {"type": "text", "content": "   ", "styles": ["underline", "strikethrough"]},
        {"type": "text", "content": "|"},
        {"type": "text", "content": "  emphasis  ", "styles": ["emphasis"]},
        {"type": "text", "content": "|"},
        {"type": "text", "content": "  bold  ", "styles": ["bold"]},
        {"type": "text", "content": "|"},
        {"type": "text", "content": "a b", "styles": ["underline"]},
    ]
    middle = _middle(_page(0, TextBlock(type="text", index=0, content=content)))
    original = deepcopy(middle)

    result = render_docx(middle)
    document = Document(BytesIO(result))
    document_xml = _part(result, "word/document.xml")

    assert document.paragraphs[0].text == (
        "\u00a0\u00a0left|right\u00a0\u00a0|\u00a0\u00a0\u00a0|\u00a0\u00a0emphasis\u00a0\u00a0|  bold  |a b"
    )
    assert document.paragraphs[0].text.count("\u00a0") == 11
    assert '<w:em w:val="underDot"' in document_xml
    assert '<w:em w:val="dot"' not in document_xml
    assert middle == original


def test_bookmark_names_are_sanitized_and_collision_safe() -> None:
    """验证不同原 anchor 清洗后冲突时仍生成唯一 Word bookmark。"""
    first = ParagraphTitleBlock(
        type="paragraph_title",
        index=0,
        level=2,
        anchor="section-a",
        content=_inline("First"),
    )
    second = ParagraphTitleBlock(
        type="paragraph_title",
        index=1,
        level=2,
        anchor="section a",
        content=_inline("Second"),
    )

    document_xml = _part(render_docx(_middle(_page(0, first, second))), "word/document.xml")

    assert 'w:name="section_a"' in document_xml
    assert document_xml.count("<w:bookmarkStart") == 2
    names = document_xml.split('w:bookmarkStart w:id="')[1:]
    assert len({value.split('w:name="', 1)[1].split('"', 1)[0] for value in names}) == 2


def test_index_only_anchor_falls_back_to_plain_text_without_dangling_link() -> None:
    """验证目录独有 anchor 不会被误注册为没有正文 bookmark 的内部链接。"""
    index = IndexBlock(
        type="index",
        index=0,
        content=[
            ParagraphTitleBlock(
                type="paragraph_title",
                level=2,
                anchor="missing-target",
                content=_inline("Missing target\t8"),
            )
        ],
    )

    document_xml = _part(render_docx(_middle(_page(0, index))), "word/document.xml")

    assert "Missing target" in document_xml
    assert "w:anchor=" not in document_xml
    assert "w:bookmarkStart" not in document_xml


def test_docx_uses_default_planner_without_source_page_boundaries() -> None:
    """验证固定默认 DOCX 隐藏辅助块、合并续段且不写源页硬分页。"""
    middle = _middle(
        _page(
            0,
            TextBlock(type="text", index=0, content=_inline("international")),
            PageAuxTextBlock(type="header", index=1, content=_inline("HEADER")),
        ),
        _page(1, TextBlock(type="text", index=0, content=_inline("continuation"), continues_prev=True)),
    )

    result = render_docx(middle)
    document = Document(BytesIO(result))

    assert [paragraph.text for paragraph in document.paragraphs] == ["international continuation"]
    assert "HEADER" not in [paragraph.text for paragraph in document.paragraphs]
    assert 'w:br w:type="page"' not in _part(result, "word/document.xml")


def test_list_preserves_markers_and_uses_hanging_indents_without_numbering() -> None:
    """验证列表不重建 numbering.xml，只保留 marker 和递归缩进。"""
    nested = ListBlock(
        type="list",
        content=[TextBlock(type="text", content=_inline("- nested item"))],
    )
    block = ListBlock(
        type="list",
        index=0,
        content=[TextBlock(type="text", content=_inline("1. first item")), nested],
    )

    result = render_docx(_middle(_page(0, block)))
    document = Document(BytesIO(result))
    document_xml = _part(result, "word/document.xml")

    assert [paragraph.text for paragraph in document.paragraphs] == ["1. first item", "- nested item"]
    assert abs(int(document.paragraphs[0].paragraph_format.left_indent) - int(Mm(6))) < 635
    assert abs(int(document.paragraphs[0].paragraph_format.first_line_indent) + int(Mm(6))) < 635
    assert abs(int(document.paragraphs[1].paragraph_format.left_indent) - int(Mm(12))) < 635
    assert "<w:numPr" not in document_xml


def test_display_formula_tag_uses_center_and_right_tabs() -> None:
    """验证公式主体为 OMML，tag 通过右对齐 tab 单独输出。"""
    block = EquationBlock(type="equation", index=0, content=r"x^2=1\tag{9}")

    result = render_docx(_middle(_page(0, block)))
    document_xml = _part(result, "word/document.xml")
    document = Document(BytesIO(result))

    assert "<m:oMath" in document_xml
    assert "<m:oMathPara" not in document_xml
    assert document.paragraphs[0].text == "\t\t(9)"
    assert document_xml.count("<w:tab/>") == 2
    assert 'w:val="center"' in document_xml
    assert 'w:val="right"' in document_xml


def test_display_formula_without_tag_uses_math_paragraph() -> None:
    """验证无编号块公式使用居中的 m:oMathPara。"""
    result = render_docx(_middle(_page(0, EquationBlock(type="equation", index=0, content=r"x^2=1"))))
    document_xml = _part(result, "word/document.xml")
    root = etree.fromstring(document_xml.encode("utf-8"))
    namespace = {
        "m": "http://schemas.openxmlformats.org/officeDocument/2006/math",
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
    }
    math_paragraph = root.find(".//m:oMathPara", namespaces=namespace)

    assert "<m:oMathPara" in document_xml
    assert '<m:jc m:val="center"' in document_xml
    assert math_paragraph is not None
    assert math_paragraph.getparent().tag == f"{{{namespace['w']}}}p"


def test_formula_conversion_failure_falls_back_to_image_then_visible_latex() -> None:
    """验证块公式优先图片回退，行内公式保留可见 LaTeX。"""
    unsupported = r"\frac{a"
    middle = _middle(
        _page(
            0,
            EquationBlock(
                type="equation",
                index=0,
                content=unsupported,
                image_base64=_png_uri(),
            ),
            TextBlock(type="text", index=1, content=_inline(f"before <eq>{unsupported}</eq> after")),
        )
    )

    result = render_docx(middle)
    document = Document(BytesIO(result))
    document_xml = _part(result, "word/document.xml")

    assert len(document.inline_shapes) == 1
    assert unsupported in "".join(paragraph.text for paragraph in document.paragraphs)
    assert "<a:blip" in document_xml


def test_genfrac_formula_renders_native_omml_in_title_and_list() -> None:
    """验证正文标题和列表中的规范 genfrac 均输出原生双行 OMML。"""
    formula = (
        r"\left(x+a\right)^{n}=\sum_{k=0}^{n}"
        r"\left(\genfrac{}{}{0pt}{}{n}{k}\right)x^{k}a^{n-k}"
    )
    title = ParagraphTitleBlock(
        type="paragraph_title",
        index=0,
        level=2,
        content=[{"type": "text", "content": "Title "}, {"type": "equation_inline", "content": formula}],
    )
    list_block = ListBlock(
        type="list",
        index=1,
        content=[
            TextBlock(
                type="text",
                content=[
                    {"type": "text", "content": "1. before "},
                    {"type": "equation_inline", "content": formula},
                    {"type": "text", "content": " after"},
                ],
            )
        ],
    )

    result = render_docx(_middle(_page(0, title, list_block)))
    document_xml = _part(result, "word/document.xml")

    assert document_xml.count("<m:oMath") == 2
    assert document_xml.count("<m:m>") == 2
    assert document_xml.count("<m:d>") >= 2
    assert r"\genfrac" not in document_xml


def test_inline_bare_scripts_use_word_superscript_runs_without_placeholder_boxes() -> None:
    """验证无底数公式片段改用上下标 run，避免 OMML 可见占位框。"""
    block = TextBlock(
        type="text",
        index=0,
        content=[
            {"type": "text", "content": "S K"},
            {"type": "equation_inline", "content": "^{1/2}"},
            {"type": "text", "content": "/cm"},
            {"type": "equation_inline", "content": "^{-1}"},
        ],
    )

    result = render_docx(_middle(_page(0, block)))
    document_xml = _part(result, "word/document.xml")

    assert Document(BytesIO(result)).paragraphs[0].text == "S K1/2/cm-1"
    assert document_xml.count('w:vertAlign w:val="superscript"') == 2
    assert "<m:oMath" not in document_xml


def test_required_image_error_contains_public_block_location() -> None:
    """验证缺少 resolver 的必需图片通过公共异常暴露完整定位。"""
    image = ImageBlock(
        type="image",
        index=3,
        content=[
            ImageBodyBlock(
                type="image_body",
                index=3,
                content="description",
                image_path="images/missing.png",
            )
        ],
    )

    with pytest.raises(DocxRenderError, match="asset_resolver") as exc_info:
        render_docx(_middle(_page(7, image)))

    assert exc_info.value.page_idx == 7
    assert exc_info.value.block_index == 3
    assert exc_info.value.block_type == "image"


def test_image_alt_text_and_visual_child_order_are_preserved() -> None:
    """验证图片 alt description 及 caption/body/footnote 源顺序。"""
    image = ImageBlock.model_validate(
        {
            "type": "image",
            "index": 0,
            "content": [
                {"type": "image_caption", "content": _inline("before")},
                {
                    "type": "image_body",
                    "index": 0,
                    "content": "<p>diagram description</p>",
                    "image_base64": _png_uri(),
                },
                {"type": "image_footnote", "content": _inline("after")},
            ],
        }
    )

    document_xml = _part(render_docx(_middle(_page(0, image))), "word/document.xml")

    assert document_xml.index("before") < document_xml.index("diagram description")
    assert document_xml.index("diagram description") < document_xml.index("after")
    assert 'descr="diagram description"' in document_xml


def test_mineru_svg_writes_native_svg_relationship_with_png_fallback() -> None:
    """验证 DOCX DrawingML 同时引用原生 SVG 与高密度 PNG fallback。"""
    image = ImageBlock.model_validate(
        {
            "type": "image",
            "index": 0,
            "content": [
                {
                    "type": "image_body",
                    "index": 0,
                    "content": "vector formula",
                    "image_base64": _generated_svg_uri(),
                }
            ],
        }
    )

    result = render_docx(_middle(_page(0, image)))
    with zipfile.ZipFile(BytesIO(result)) as archive:
        media = sorted(name for name in archive.namelist() if name.startswith("word/media/"))
        document_xml = etree.fromstring(archive.read("word/document.xml"))
        relationships = etree.fromstring(archive.read("word/_rels/document.xml.rels"))
        content_types = archive.read("[Content_Types].xml").decode("utf-8")

    assert any(name.endswith(".png") for name in media)
    assert any(name.endswith(".svg") for name in media)
    namespace = {
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "asvg": "http://schemas.microsoft.com/office/drawing/2016/SVG/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
    }
    fallback_id = document_xml.xpath("string(.//a:blip/@r:embed)", namespaces=namespace)
    svg_id = document_xml.xpath("string(.//asvg:svgBlip/@r:embed)", namespaces=namespace)
    targets = {relationship.get("Id"): relationship.get("Target") for relationship in relationships}
    assert fallback_id != svg_id
    assert targets[fallback_id].endswith(".png")
    assert targets[svg_id].endswith(".svg")
    assert "image/svg+xml" in content_types


def test_html_table_materializes_merges_inline_content_link_and_image() -> None:
    """验证复杂 HTML 表格生成合并几何，并复用单元格行内与素材 visitor。"""
    html = (
        "<table><thead><tr><th rowspan='2'>A</th><th colspan='2'>B</th></tr></thead>"
        "<tbody><tr><td><strong>x</strong><eq>y</eq>"
        "<a href='https://example.com/table'>link</a></td>"
        f"<td><img src='{_png_uri()}' alt='cell image'/></td></tr></tbody></table>"
    )
    table = TableBlock(
        type="table",
        index=0,
        content=[TableBodyBlock(type="table_body", index=0, content=html)],
    )

    result = render_docx(_middle(_page(0, table)))
    document_xml = _part(result, "word/document.xml")
    relationships = _part(result, "word/_rels/document.xml.rels")
    document = Document(BytesIO(result))

    assert len(document.tables) == 1
    assert '<w:gridSpan w:val="2"' in document_xml
    assert "<w:vMerge" in document_xml
    assert "<w:tblHeader" in document_xml
    assert "<m:oMath" in document_xml
    assert "<a:blip" in document_xml
    assert "https://example.com/table" in relationships


def test_html_table_does_not_restore_legacy_text_style_tags() -> None:
    """验证旧 text-style 容器只保留文字，标准 strong 标签继续恢复粗体。"""
    table = TableBlock(
        type="table",
        index=0,
        content=[
            TableBodyBlock(
                type="table_body",
                index=0,
                content=('<table><tr><td><text style="bold">legacy</text><strong>standard</strong></td></tr></table>'),
            )
        ],
    )

    document = Document(BytesIO(render_docx(_middle(_page(0, table)))))
    runs = [run for paragraph in document.tables[0].cell(0, 0).paragraphs for run in paragraph.runs if run.text]

    assert [(run.text, bool(run.bold)) for run in runs] == [("legacy", False), ("standard", True)]


def test_html_table_cell_lists_keep_item_boundaries() -> None:
    """验证 HTML 单元格中的有序和无序列表不会串成连续文本。"""
    html = (
        "<table><tr><td>Items:<ul><li>first</li><li>second</li></ul>"
        "<ol start='3'><li>third</li><li>fourth</li></ol></td></tr></table>"
    )
    table = TableBlock(
        type="table",
        index=0,
        content=[TableBodyBlock(type="table_body", index=0, content=html)],
    )

    document = Document(BytesIO(render_docx(_middle(_page(0, table)))))

    assert document.tables[0].cell(0, 0).text == "Items:\n- first\n- second\n3. third\n4. fourth"


def test_html_table_direct_nested_table_is_materialized_recursively() -> None:
    """验证单元格直接子 table 通过绑定 writer 递归物化，而不是压成纯文本。"""
    html = "<table><tr><td>before<table><tr><td>nested</td></tr></table>after</td></tr></table>"
    table = TableBlock(
        type="table",
        index=0,
        content=[TableBodyBlock(type="table_body", index=0, content=html)],
    )

    result = render_docx(_middle(_page(0, table)))
    document_xml = _part(result, "word/document.xml")
    cell = Document(BytesIO(result)).tables[0].cell(0, 0)

    assert document_xml.count("<w:tbl>") == 2
    assert "before" in document_xml and "nested" in document_xml and "after" in document_xml
    assert len(cell.tables) == 1
    assert [paragraph.text for paragraph in cell.paragraphs if paragraph.text] == ["before", "after"]


def test_html_table_nested_inside_wrapper_keeps_table_and_source_order() -> None:
    """验证 div 等包装层中的嵌套表格不会被压平成连续文本。"""
    html = "<table><tr><td><div>before<table><tr><td>nested</td></tr></table>after</div></td></tr></table>"
    table = TableBlock(
        type="table",
        index=0,
        content=[TableBodyBlock(type="table_body", index=0, content=html)],
    )

    result = render_docx(_middle(_page(0, table)))
    cell = Document(BytesIO(result)).tables[0].cell(0, 0)

    assert len(cell.tables) == 1
    assert [paragraph.text for paragraph in cell.paragraphs if paragraph.text] == ["before", "after"]
    assert _part(result, "word/document.xml").count("<w:tbl>") == 2


def test_html_table_fallback_rolls_back_partial_table_and_relationships() -> None:
    """验证单元格素材失败时移除半成品表格和其新增关系，再只写整体图片。"""
    html = f"<table><tr><td><img src='{_png_uri(size=(8, 8))}'/></td><td><img src='images/missing.png'/></td></tr></table>"
    table = TableBlock(
        type="table",
        index=0,
        content=[
            TableBodyBlock(
                type="table_body",
                index=0,
                content=html,
                image_base64=_png_uri(size=(10, 10)),
            )
        ],
    )

    result = render_docx(_middle(_page(0, table)))
    document = Document(BytesIO(result))
    with zipfile.ZipFile(BytesIO(result)) as archive:
        media = [name for name in archive.namelist() if name.startswith("word/media/")]

    assert len(document.tables) == 0
    assert len(document.inline_shapes) == 1
    assert len(media) == 1


def test_html_table_remote_cell_image_uses_safe_link_fallback() -> None:
    """验证远程单元格图片输出可点击 alt，而不是中断整份 DOCX。"""
    html = '<table><tr><td><img src="https://example.com/logo.png" alt="Logo"></td></tr></table>'
    table = TableBlock(
        type="table",
        index=0,
        content=[TableBodyBlock(type="table_body", index=0, content=html)],
    )

    result = render_docx(_middle(_page(0, table)))
    document = Document(BytesIO(result))

    assert document.tables[0].cell(0, 0).text == "Logo"
    assert len(document.inline_shapes) == 0
    assert "https://example.com/logo.png" in _part(result, "word/_rels/document.xml.rels")


def test_html_table_cell_image_is_limited_to_merged_cell_width() -> None:
    """验证窄列图片宽度不超过扣除单元格左右内边距后的 tcW。"""
    cells = [f"<td>{'<img src=' + repr(_png_uri(size=(300, 100))) + '/>' if index == 0 else index}</td>" for index in range(10)]
    html = f"<table><tr>{''.join(cells)}</tr></table>"
    table = TableBlock(
        type="table",
        index=0,
        content=[TableBodyBlock(type="table_body", index=0, content=html)],
    )

    document = Document(BytesIO(render_docx(_middle(_page(0, table)))))
    cell = document.tables[0].cell(0, 0)
    cell_width_twips = int(cell._tc.get_or_add_tcPr().get_or_add_tcW().w)
    available_width_emu = int(Twips(cell_width_twips - 240))

    assert len(document.inline_shapes) == 1
    assert int(document.inline_shapes[0].width) <= available_width_emu


def test_invalid_html_table_falls_back_to_table_image() -> None:
    """验证 HTML 占位网格非法时使用表格图片，不留下损坏表格。"""
    table = TableBlock(
        type="table",
        index=0,
        content=[
            TableBodyBlock(
                type="table_body",
                index=0,
                content="<table><tr><td rowspan='2'>A</td></tr></table>",
                image_base64=_png_uri(),
            )
        ],
    )

    result = render_docx(_middle(_page(0, table)))
    document = Document(BytesIO(result))

    assert len(document.tables) == 0
    assert len(document.inline_shapes) == 1

    without_fallback = TableBlock(
        type="table",
        index=0,
        content=[
            TableBodyBlock(
                type="table_body",
                index=0,
                content="<table><tr><td rowspan='2'>A</td></tr></table>",
            )
        ],
    )
    with pytest.raises(DocxRenderError, match="no image fallback"):
        render_docx(_middle(_page(0, without_fallback)))


@pytest.mark.parametrize(
    "image_payload",
    [
        pytest.param({}, id="without-image"),
        pytest.param({"image_base64": _png_uri()}, id="base64-image"),
        pytest.param({"image_path": "images/table.png"}, id="sidecar-image"),
    ],
)
def test_spatial_table_preserves_preformatted_text_without_assets(image_payload: dict[str, str]) -> None:
    """验证空间表格原样保留排版字符，并完全忽略可选图片素材。"""
    content = "  A    B\t说明\n\n1    2\t中\x01文\n" + "X" * 240
    spatial = TableBlock(
        type="table",
        index=0,
        content=[
            TableBodyBlock(
                type="table_body",
                index=0,
                content=content,
                **image_payload,
            )
        ],
    )
    resolver = Mock(side_effect=AssertionError("空间表格不应解析图片"))

    result = render_docx(_middle(_page(0, spatial)), asset_resolver=resolver)
    document = Document(BytesIO(result))
    document_xml = _part(result, "word/document.xml")
    relationships = _part(result, "word/_rels/document.xml.rels")

    resolver.assert_not_called()
    assert document.paragraphs[0].text == content.replace("\x01", "\ufffd")
    assert document.paragraphs[0].style.name == "MinerU Spatial Table"
    assert document.styles["MinerU Spatial Table"].font.name == document.styles["MinerU Code"].font.name
    assert document.styles["MinerU Spatial Table"].font.size == document.styles["MinerU Code"].font.size
    assert document.styles["MinerU Spatial Table"].paragraph_format.line_spacing == 1.0
    assert 'w:pStyle w:val="MinerUSpatialTable"' in document_xml
    assert 'xml:space="preserve"' in document_xml
    assert "<w:tab/>" in document_xml
    assert document_xml.count("<w:br/>") == 3
    assert len(document.inline_shapes) == 0
    assert "relationships/image" not in relationships
    with zipfile.ZipFile(BytesIO(result)) as archive:
        assert not any(name.startswith("word/media/") for name in archive.namelist())


@pytest.mark.parametrize("content", ["", " \n\t"])
def test_spatial_table_without_text_uses_preferred_sidecar_image(content: str) -> None:
    """验证空或纯空白空间表格遵循 image_path 优先的公共图片契约。"""
    spatial = TableBlock(
        type="table",
        index=7,
        content=[
            TableBodyBlock(
                type="table_body",
                index=7,
                content=content,
                image_base64=_png_uri(),
                image_path="images/unused.png",
            )
        ],
    )
    resolver = Mock(return_value=_png_bytes(size=(16, 10)))

    document = Document(BytesIO(render_docx(_middle(_page(3, spatial)), asset_resolver=resolver)))

    resolver.assert_called_once_with("images/unused.png")
    assert len(document.inline_shapes) == 1


def test_spatial_table_without_text_uses_sidecar_resolver() -> None:
    """验证空空间表格的相对图片路径只通过注入的 resolver 加载。"""
    spatial = TableBlock(
        type="table",
        index=7,
        content=[
            TableBodyBlock(
                type="table_body",
                index=7,
                content="",
                image_path="images/table.png",
            )
        ],
    )
    resolver = Mock(return_value=_png_bytes())

    document = Document(BytesIO(render_docx(_middle(_page(3, spatial)), asset_resolver=resolver)))

    resolver.assert_called_once_with("images/table.png")
    assert len(document.inline_shapes) == 1


def test_remote_only_equation_table_and_chart_use_docx_link_fallbacks() -> None:
    """验证所有远程-only 图片载荷都会进入 DOCX 可点击链接回退。"""
    middle = _middle(
        _page(
            0,
            EquationBlock(
                type="equation",
                index=0,
                content="",
                image_url="https://example.com/formula.png",
            ),
            TableBlock(
                type="table",
                index=1,
                content=[
                    TableBodyBlock(
                        type="table_body",
                        index=1,
                        content="",
                        image_url="https://example.com/table.png",
                    )
                ],
            ),
            ChartBlock(
                type="chart",
                index=2,
                sub_type="bar",
                content=[
                    ChartBodyBlock(
                        type="chart_body",
                        index=2,
                        content="",
                        image_url="https://example.com/chart.png",
                    )
                ],
            ),
        )
    )

    result = render_docx(middle)
    document = Document(BytesIO(result))
    relationships = _part(result, "word/_rels/document.xml.rels")

    assert [paragraph.text for paragraph in document.paragraphs] == ["formula", "table", "bar"]
    assert all(
        target in relationships
        for target in (
            "https://example.com/formula.png",
            "https://example.com/table.png",
            "https://example.com/chart.png",
        )
    )


@pytest.mark.parametrize("content", ["", " \n\t"])
def test_spatial_table_without_text_or_image_raises_contextual_error(content: str) -> None:
    """验证空间表格既无有效文本也无图片时抛出带父表格定位的异常。"""
    spatial = TableBlock(
        type="table",
        index=7,
        content=[TableBodyBlock(type="table_body", index=7, content=content)],
    )

    with pytest.raises(DocxRenderError, match="does not contain text content or image") as exc_info:
        render_docx(_middle(_page(3, spatial)))

    assert exc_info.value.page_idx == 3
    assert exc_info.value.block_index == 7
    assert exc_info.value.block_type == "table"


def test_spatial_table_without_text_rejects_invalid_fallback_image() -> None:
    """验证空空间表格的损坏图片不会被静默忽略。"""
    spatial = TableBlock(
        type="table",
        index=7,
        content=[
            TableBodyBlock(
                type="table_body",
                index=7,
                content="",
                image_base64="data:image/png;base64,bm90LWEtcG5n",
            )
        ],
    )

    with pytest.raises(DocxRenderError) as exc_info:
        render_docx(_middle(_page(3, spatial)))

    assert exc_info.value.page_idx == 3
    assert exc_info.value.block_index == 7
    assert exc_info.value.block_type == "table"


def test_spatial_table_preserves_caption_body_footnote_order() -> None:
    """验证空间表格正文仍严格位于原始 caption 与 footnote 之间。"""
    spatial = TableBlock(
        type="table",
        index=1,
        content=[
            TableAnnotationBlock(type="table_caption", index=0, content=_inline("Table caption")),
            TableBodyBlock(type="table_body", index=1, content="  A    B\n  1    2"),
            TableAnnotationBlock(type="table_footnote", index=2, content=_inline("Table footnote")),
        ],
    )

    document = Document(BytesIO(render_docx(_middle(_page(0, spatial)))))

    assert [paragraph.text for paragraph in document.paragraphs] == [
        "Table caption",
        "  A    B\n  1    2",
        "Table footnote",
    ]
    assert [paragraph.style.name for paragraph in document.paragraphs] == [
        "MinerU Caption",
        "MinerU Spatial Table",
        "MinerU Footnote",
    ]


def test_chart_renders_image_then_structured_html_table() -> None:
    """验证图表图片后继续输出可编辑的 HTML 数据表。"""
    chart = ChartBlock(
        type="chart",
        index=0,
        content=[
            ChartBodyBlock(
                type="chart_body",
                index=0,
                content="<table><tr><th>A</th></tr><tr><td>1</td></tr></table>",
                image_base64=_png_uri(),
            )
        ],
    )

    result = render_docx(_middle(_page(0, chart)))
    document = Document(BytesIO(result))
    document_xml = _part(result, "word/document.xml")

    assert len(document.inline_shapes) == 1
    assert len(document.tables) == 1
    assert document_xml.index("<a:blip") < document_xml.index("<w:tbl>")


def test_chart_invalid_html_keeps_successful_image_fallback() -> None:
    """验证图表结构表损坏时保留已成功写入的图片，而不是终止文档。"""
    chart = ChartBlock(
        type="chart",
        index=0,
        content=[
            ChartBodyBlock(
                type="chart_body",
                index=0,
                content="<table><tr><td rowspan='2'>broken</td></tr></table>",
                image_base64=_png_uri(),
            )
        ],
    )

    document = Document(BytesIO(render_docx(_middle(_page(0, chart)))))

    assert len(document.inline_shapes) == 1
    assert len(document.tables) == 0


def test_code_and_algorithm_keep_line_breaks_styles_and_inline_math() -> None:
    """验证代码换行与算法行内公式、上下标均保留。"""
    code = CodeBlock.model_validate(
        {
            "type": "code",
            "index": 0,
            "sub_type": "code",
            "guess_lang": "python",
            "content": [{"type": "code_body", "index": 0, "content": "a = 1\nb = 2"}],
        }
    )
    algorithm = CodeBlock.model_validate(
        {
            "type": "code",
            "index": 1,
            "sub_type": "algorithm",
            "content": [
                {
                    "type": "algorithm_body",
                    "index": 1,
                    "content": [
                        {"type": "text", "content": "T"},
                        {"type": "text", "content": "n", "styles": ["subscript"]},
                        {"type": "text", "content": " = "},
                        {"type": "equation_inline", "content": "x^2"},
                        {"type": "text", "content": "\nnext"},
                    ],
                }
            ],
        }
    )

    result = render_docx(_middle(_page(0, code, algorithm)))
    document = Document(BytesIO(result))
    document_xml = _part(result, "word/document.xml")

    assert [paragraph.style.name for paragraph in document.paragraphs] == ["MinerU Code", "MinerU Code"]
    assert "a = 1\nb = 2" == document.paragraphs[0].text
    assert "<m:oMath" in document_xml
    assert 'w:vertAlign w:val="subscript"' in document_xml
