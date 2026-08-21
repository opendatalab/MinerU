from __future__ import annotations

from collections.abc import Callable

import pytest
from bs4.element import Tag
from docx import Document
from docx.oxml.ns import qn
from docx.table import _Cell

from mineru.render._internal.docx.table import (
    DocxTableError,
    NestedTableWriter,
    materialize_docx_table,
    materialize_docx_tables,
    parse_html_table,
    parse_html_tables,
)


def _text_filler(
    observed: list[tuple[str, str]],
) -> Callable[[_Cell, Tag, NestedTableWriter], None]:
    """构造记录原始 Tag 并把直接文本写入 Word 单元格的测试回调。"""

    def fill_cell(cell: _Cell, source: Tag, write_nested: NestedTableWriter) -> None:
        """记录回调来源，并忽略本用例不涉及的嵌套写入器。"""
        del write_nested
        observed.append((source.name, source.get_text(" ", strip=True)))
        cell.paragraphs[0].add_run(source.get_text(" ", strip=True))

    return fill_cell


def _xml_value(element: object, attribute: str) -> str | None:
    """读取测试 XML 元素中的 WordprocessingML 属性。"""
    return element.get(qn(attribute))  # type: ignore[union-attr]


def test_parse_and_materialize_simple_table_with_header_tags() -> None:
    """验证简单表格、th 回调、Table Grid 和重复表头标记。"""
    html = (
        '<table><thead><tr><th data-id="h1">Name</th><th data-id="h2">Value</th></tr></thead>'
        '<tbody><tr><td data-id="c1">A</td><td data-id="c2">1</td></tr></tbody></table>'
    )
    grid = parse_html_tables(html)[0]

    assert (grid.row_count, grid.column_count) == (2, 2)
    assert grid.header_rows == (0,)
    assert [(cell.row, cell.column, cell.is_header) for cell in grid.cells] == [
        (0, 0, True),
        (0, 1, True),
        (1, 0, False),
        (1, 1, False),
    ]

    document = Document()
    observed: list[tuple[str, str]] = []
    table = materialize_docx_table(
        document,
        grid,
        width_twips=6000,
        fill_cell=_text_filler(observed),
    )

    assert observed == [("th", "Name"), ("th", "Value"), ("td", "A"), ("td", "1")]
    assert table.style.name == "Table Grid"
    assert table.cell(0, 0).text == "Name"
    assert table.cell(1, 1).text == "1"
    assert len(table.rows[0]._tr.xpath("./w:trPr/w:tblHeader")) == 1


def test_materialize_rowspan_and_colspan_as_gridspan_and_vmerge() -> None:
    """验证二维合并写出一致的 gridSpan、vMerge 和 origin 回调。"""
    html = '<table><tr><th rowspan="2" colspan="2">A</th><th>B</th></tr><tr><td>C</td></tr></table>'
    document = Document()
    observed: list[tuple[str, str]] = []

    table = materialize_docx_tables(
        document,
        html,
        width_twips=6000,
        fill_cell=_text_filler(observed),
    )[0]

    assert observed == [("th", "A"), ("th", "B"), ("td", "C")]
    first_row_cells = table._tbl.tr_lst[0].tc_lst
    second_row_cells = table._tbl.tr_lst[1].tc_lst
    assert _xml_value(first_row_cells[0].tcPr.gridSpan, "w:val") == "2"
    assert _xml_value(first_row_cells[0].tcPr.vMerge, "w:val") == "restart"
    assert _xml_value(second_row_cells[0].tcPr.gridSpan, "w:val") == "2"
    assert _xml_value(second_row_cells[0].tcPr.vMerge, "w:val") is None
    assert first_row_cells[0].tcPr.tcW.w == 4000
    assert second_row_cells[0].tcPr.tcW.w == 4000


def test_parse_and_materialize_multiple_top_level_tables() -> None:
    """验证同一 HTML 片段中的多个顶层表格按源码顺序独立创建。"""
    html = '<table data-id="one"><tr><td>A</td></tr></table><div><table data-id="two"><tr><td>B</td></tr></table></div>'
    grids = parse_html_tables(html)
    assert [grid.tag.get("data-id") for grid in grids] == ["one", "two"]

    document = Document()
    observed: list[tuple[str, str]] = []
    tables = materialize_docx_tables(document, html, fill_cell=_text_filler(observed))

    assert len(tables) == 2
    assert len(document.tables) == 2
    assert [table.cell(0, 0).text for table in tables] == ["A", "B"]


@pytest.mark.parametrize(
    "html",
    [
        "<table></table>",
        "<table><tr></tr></table>",
    ],
)
def test_parse_rejects_table_without_cells(html: str) -> None:
    """验证无行或无任何单元格的表格被安全拒绝。"""
    with pytest.raises(DocxTableError):
        parse_html_tables(html)


@pytest.mark.parametrize(
    ("attribute", "value"),
    [
        ("rowspan", "0"),
        ("rowspan", "-1"),
        ("rowspan", "1.5"),
        ("rowspan", "x"),
        ("colspan", "0"),
        ("colspan", ""),
    ],
)
def test_parse_rejects_invalid_spans(attribute: str, value: str) -> None:
    """验证 rowspan/colspan 只接受严格正整数。"""
    html = f'<table><tr><td {attribute}="{value}">A</td></tr></table>'

    with pytest.raises(DocxTableError, match=f"Invalid {attribute}"):
        parse_html_tables(html)


def test_parse_rejects_rowspan_out_of_bounds() -> None:
    """验证跨出最后一行的 rowspan 被判定为越界。"""
    with pytest.raises(DocxTableError, match="rowspan exceeds table bounds"):
        parse_html_tables('<table><tr><td rowspan="2">A</td></tr></table>')


def test_parse_rejects_span_overlap_after_skipping_occupied_coordinate() -> None:
    """验证跳过行首占位后，colspan 仍不得覆盖后方 rowspan 占位。"""
    html = '<table><tr><td rowspan="2">A</td><td>B</td><td rowspan="2">C</td></tr><tr><td colspan="2">D</td></tr></table>'

    with pytest.raises(DocxTableError, match="overlaps an occupied coordinate"):
        parse_html_tables(html)


def test_parse_rejects_non_rectangular_occupancy() -> None:
    """验证逻辑行宽不一致且没有 rowspan 补位时拒绝非矩形表格。"""
    html = "<table><tr><td>A</td><td>B</td></tr><tr><td>C</td></tr></table>"

    with pytest.raises(DocxTableError, match="occupancy must be rectangular"):
        parse_html_tables(html)


def test_materialized_geometry_uses_exact_dxa_widths_without_fixed_row_height() -> None:
    """验证 tblW、tblGrid、每行 tcW 的总宽完全一致且不存在固定行高。"""
    html = '<table><tr><td colspan="2">A</td><td>B</td></tr><tr><td>C</td><td>D</td><td>E</td></tr></table>'
    document = Document()
    table = materialize_docx_tables(
        document,
        html,
        width_twips=10001,
        fill_cell=_text_filler([]),
    )[0]

    table_width = table._tbl.tblPr.find(qn("w:tblW"))
    table_layout = table._tbl.tblPr.find(qn("w:tblLayout"))
    assert table_width is not None
    assert (_xml_value(table_width, "w:type"), _xml_value(table_width, "w:w")) == (
        "dxa",
        "10001",
    )
    assert table_layout is not None
    assert _xml_value(table_layout, "w:type") == "fixed"

    grid_widths = [column.w.twips for column in table._tbl.tblGrid.gridCol_lst]
    assert grid_widths == [3334, 3334, 3333]
    for row in table._tbl.tr_lst:
        assert sum(cell.tcPr.tcW.w for cell in row.tc_lst) == 10001
        assert row.xpath("./w:trPr/w:trHeight") == []


def test_callback_can_materialize_four_nested_table_levels() -> None:
    """验证回调可使用绑定写入器递归创建最多四层嵌套表格。"""
    html = _nested_table_html(4)
    document = Document()
    seen_levels: list[str] = []

    def fill_cell(cell: _Cell, source: Tag, write_nested: NestedTableWriter) -> None:
        """记录当前层级，并递归写入当前单元格中的直接子表。"""
        parent_table = source.find_parent("table")
        assert parent_table is not None
        seen_levels.append(str(parent_table.get("data-level")))
        for nested in source.find_all("table"):
            if nested.find_parent("table") is parent_table:
                write_nested(nested)
        cell.paragraphs[0].add_run(f"L{parent_table.get('data-level')}")

    materialize_docx_tables(document, html, width_twips=6000, fill_cell=fill_cell)

    assert seen_levels == ["1", "2", "3", "4"]
    assert len(document.element.body.xpath(".//w:tbl")) == 4


def test_callback_rejects_fifth_nested_table_level() -> None:
    """验证第五层嵌套表格在回调尝试写入时抛出明确错误。"""
    document = Document()

    def fill_cell(cell: _Cell, source: Tag, write_nested: NestedTableWriter) -> None:
        """只递归写入直接子表，以触发统一深度保护。"""
        del cell
        parent_table = source.find_parent("table")
        assert parent_table is not None
        for nested in source.find_all("table"):
            if nested.find_parent("table") is parent_table:
                write_nested(nested)

    with pytest.raises(DocxTableError, match="depth exceeds 4"):
        materialize_docx_tables(
            document,
            _nested_table_html(5),
            width_twips=6000,
            fill_cell=fill_cell,
        )


def test_parse_single_tag_preserves_the_original_cell_tag_objects() -> None:
    """验证中性解析器保留调用方提供的原始 td/th Tag，而不是复制文本。"""
    soup_grid = parse_html_tables('<table><tr><th data-key="a">A</th><td data-key="b">B</td></tr></table>')[0]
    table_tag = soup_grid.tag
    parsed_again = parse_html_table(table_tag)
    source_cells = table_tag.find_all(("th", "td"), recursive=True)

    assert [cell.tag is source for cell, source in zip(parsed_again.cells, source_cells)] == [True, True]
    assert [cell.tag.get("data-key") for cell in parsed_again.cells] == ["a", "b"]


def _nested_table_html(depth: int, level: int = 1) -> str:
    """生成指定总层数的单单元格嵌套表格 HTML。"""
    nested = _nested_table_html(depth - 1, level + 1) if depth > 1 else ""
    return f'<table data-level="{level}"><tr><td>{nested}</td></tr></table>'
